"""
Document export functionality for the Whisper Transcription App.
Handles exporting transcription results to various document formats.
"""

import io
from pathlib import Path
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
import config
from transcription.formatter import TranscriptionFormatter


class DocumentExporter:
    """Handles exporting transcription results to document formats."""
    
    def __init__(self):
        """Initialize the document exporter."""
        self.formatter = TranscriptionFormatter()
    
    def create_text_download(self, content: str, filename: str) -> bytes:
        """
        Create a text file for download.
        
        Args:
            content: Text content to include
            filename: Base filename (extension will be added)
            
        Returns:
            Bytes content for download
        """
        return content.encode('utf-8')
    
    def create_docx_download(self, transcription_data: Dict[str, Any], 
                           format_type: str, filename: str) -> Optional[bytes]:
        """
        Create a DOCX document for download.
        
        Args:
            transcription_data: Processed transcription data
            format_type: Type of format ('word_timestamps' or 'segment_timestamps')
            filename: Base filename
            
        Returns:
            Bytes content for download, or None if failed
        """
        try:
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Audio Transcription Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add summary information
            self._add_summary_section(doc, transcription_data)
            
            # Add main content based on format type
            if format_type == 'word_timestamps':
                self._add_word_timestamps_section(doc, transcription_data)
            elif format_type == 'segment_timestamps':
                self._add_segment_timestamps_section(doc, transcription_data)
            else:
                self._add_plain_text_section(doc, transcription_data)
            
            # Save to bytes
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            return doc_io.getvalue()
            
        except Exception as e:
            st.error(f"Error creating DOCX document: {str(e)}")
            return None
    
    def _add_summary_section(self, doc: Document, transcription_data: Dict[str, Any]):
        """Add summary information to the document."""
        doc.add_heading('Summary', level=1)
        
        summary = self.formatter.get_transcription_summary(transcription_data)
        
        table = doc.add_table(rows=len(summary), cols=2)
        table.style = 'Table Grid'
        
        for i, (key, value) in enumerate(summary.items()):
            row = table.rows[i]
            row.cells[0].text = key
            row.cells[1].text = value
            
            # Make the first column bold
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_page_break()
    
    def _add_plain_text_section(self, doc: Document, transcription_data: Dict[str, Any]):
        """Add plain text transcription to the document."""
        doc.add_heading('Transcription', level=1)
        
        text = self.formatter.format_plain_text(transcription_data)
        doc.add_paragraph(text)
    
    def _add_word_timestamps_section(self, doc: Document, transcription_data: Dict[str, Any]):
        """Add word-level timestamps to the document."""
        doc.add_heading('Word-Level Timestamps', level=1)
        
        words = transcription_data.get("words", [])
        if not words:
            doc.add_paragraph("No word-level timestamps available.")
            return
        
        current_line = []
        line_start_time = None
        
        for word_data in words:
            word = word_data.get("word", "").strip()
            start = word_data.get("start", 0)
            end = word_data.get("end", 0)
            
            if not word:
                continue
            
            if line_start_time is None:
                line_start_time = start
            
            current_line.append(word)
            
            # Break line every ~10 words or at sentence boundaries
            if (len(current_line) >= 10 or 
                word.endswith('.') or word.endswith('!') or word.endswith('?')):
                
                self._add_timestamped_paragraph(doc, current_line, line_start_time, end)
                current_line = []
                line_start_time = None
        
        # Add any remaining words
        if current_line:
            end_time = words[-1].get("end", 0)
            self._add_timestamped_paragraph(doc, current_line, line_start_time or 0, end_time)
    
    def _add_segment_timestamps_section(self, doc: Document, transcription_data: Dict[str, Any]):
        """Add segment-level timestamps to the document."""
        doc.add_heading('Segment-Level Timestamps', level=1)
        
        segments = transcription_data.get("segments", [])
        if not segments:
            doc.add_paragraph("No segment-level timestamps available.")
            return
        
        for i, segment in enumerate(segments, 1):
            start = segment.get("start", 0)
            end = segment.get("end", 0)
            text = segment.get("text", "").strip()
            
            if not text:
                continue
            
            # Add segment header
            segment_header = doc.add_heading(f'Segment {i:03d}', level=2)
            
            # Add timestamp
            start_ts = self.formatter.format_timestamp(start)
            end_ts = self.formatter.format_timestamp(end)
            timestamp_para = doc.add_paragraph()
            timestamp_run = timestamp_para.add_run(f"Time: {start_ts} → {end_ts}")
            timestamp_run.italic = True
            
            # Add text
            doc.add_paragraph(text)
            doc.add_paragraph()  # Add spacing
    
    def _add_timestamped_paragraph(self, doc: Document, words: list, start_time: float, end_time: float):
        """Add a paragraph with timestamp for word-level formatting."""
        line_text = " ".join(words)
        start_ts = self.formatter.format_timestamp(start_time)
        end_ts = self.formatter.format_timestamp(end_time)
        
        # Add timestamp
        timestamp_para = doc.add_paragraph()
        timestamp_run = timestamp_para.add_run(f"[{start_ts} → {end_ts}]")
        timestamp_run.italic = True
        timestamp_run.bold = True
        
        # Add text
        doc.add_paragraph(line_text)
        doc.add_paragraph()  # Add spacing
    
    def get_download_button_label(self, format_type: str, file_format: str) -> str:
        """
        Get appropriate label for download button.
        
        Args:
            format_type: Type of transcription format
            file_format: File format (txt or docx)
            
        Returns:
            Button label string
        """
        format_names = {
            'plain_text': 'Plain Text',
            'word_timestamps': 'Word Timestamps',
            'segment_timestamps': 'Segment Timestamps'
        }
        
        format_name = format_names.get(format_type, format_type.title())
        extension = file_format.upper()
        
        return f"📄 Download {format_name} ({extension})"
    
    def get_filename(self, base_name: str, format_type: str, file_format: str) -> str:
        """
        Generate appropriate filename for export.
        
        Args:
            base_name: Base filename
            format_type: Type of transcription format
            file_format: File format extension
            
        Returns:
            Complete filename
        """
        if not base_name:
            base_name = config.DEFAULT_FILENAME_PREFIX
        
        # Clean base name
        clean_name = "".join(c for c in base_name if c.isalnum() or c in (' ', '-', '_'))
        clean_name = clean_name.replace(' ', '_')
        
        return f"{clean_name}_{format_type}.{file_format}"
    
    def create_subtitle_download(self, transcription_data: Dict[str, Any], 
                               subtitle_format: str) -> Optional[bytes]:
        """
        Create subtitle file for download.
        
        Args:
            transcription_data: Processed transcription data
            subtitle_format: Format type ('srt' or 'vtt')
            
        Returns:
            Bytes content for download, or None if failed
        """
        try:
            if subtitle_format == 'srt':
                content = self.formatter.format_srt_subtitles(transcription_data)
            elif subtitle_format == 'vtt':
                content = self.formatter.format_vtt_subtitles(transcription_data)
            else:
                return None
            
            return content.encode('utf-8')
            
        except Exception as e:
            st.error(f"Error creating subtitle file: {str(e)}")
            return None
